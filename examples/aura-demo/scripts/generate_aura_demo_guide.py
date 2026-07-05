"""Generate AURA_Demo_Guide.docx from the current README/WALKTHROUGH and codebase facts.

Usage:
    cd backend && ..\.venv\Scripts\python.exe ../scripts/generate_aura_demo_guide.py
    # or any Python with python-docx installed
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parent.parent
DOCX_PATH = ROOT / "AURA_Demo_Guide.docx"


def set_style(doc, name, size, bold=False, color=None, space_after=Pt(6)):
    style = doc.styles[name]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)
    pformat = style.paragraph_format
    pformat.space_after = space_after
    pformat.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_para(doc, text, style="Normal", bold=False, italic=False):
    p = doc.add_paragraph(text, style=style)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold or italic:
        run = p.runs[0]
        run.bold = bold
        run.italic = italic
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.25)


def add_code(doc, lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("\n".join(lines))
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Inches(0.25)
    return p


def main():
    doc = Document()

    # Page setup
    sections = doc.sections[0]
    sections.top_margin = Inches(0.8)
    sections.bottom_margin = Inches(0.8)
    sections.left_margin = Inches(0.9)
    sections.right_margin = Inches(0.9)

    set_style(doc, "Normal", 11)
    set_style(doc, "Heading 1", 18, bold=True, color=(0x00, 0x7A, 0x3E))
    set_style(doc, "Heading 2", 14, bold=True, color=(0x00, 0x5C, 0x2E))
    set_style(doc, "Heading 3", 12, bold=True)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AURA / Assure — Phase 2 Demo Guide")
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x7A, 0x3E)
    run.font.name = "Calibri Light"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Prepared for Kevin — Financial Simplicity")
    run.font.size = Pt(14)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: 2026-06-22\nRepository: financialSimplicity/prototyping").font.size = Pt(10)
    meta.paragraph_format.space_after = Pt(24)

    doc.add_page_break()

    # 1. Executive summary
    add_heading(doc, "1. Executive Summary", level=1)
    add_para(doc,
        "AURA (Assure) is a portfolio-assurance platform that puts a deterministic rules engine at the centre "
        "and wraps it with AI comprehension, proposal and narration. Every AI action is gated by the rules "
        "engine and every action that changes state is approved by a human. The system is fully audited.")
    add_para(doc,
        "Phase 2 adds Hermes — a book-wide, self-improving remediation engine — and scales the synthetic "
        "book to 34,000 portfolios backed by a seeded market-simulation model. The demo can be run locally "
        "in minutes and is ready for deployment to Render + Vercel.")
    add_bullets(doc, [
        "34,000 synthetic portfolios in SQLite with WAL.",
        "Seeded geometric-Brownian-motion market model with virtual clock.",
        "Deterministic rules engine = source of truth; LLM is advisory only.",
        "Human-in-the-loop approval on every remediation and strategy change.",
        "Append-only audit trail.",
        "Offline MockLLM fallback when ANTHROPIC_API_KEY is absent.",
        "Muted Light Matrix institutional UI (JetBrains Mono + cool slate palette) across Command Centre, Diagnosis, Workbench and /hermes.",
    ])

    # 2. Demo philosophy
    add_heading(doc, "2. Demo Philosophy: AI Recommends, Assurance Verifies", level=1)
    add_para(doc,
        "The demo tells a single story: compliance is never delegated to the LLM. "
        "The rules engine is the LAW. Hermes strategy variables are JUDGMENT. Humans hold the keys.")
    add_bullets(doc, [
        "Mandate rules live with each portfolio and are evaluated by core/rules_engine.py only.",
        "Hermes may propose strategy tweaks in agents/hermes/strategy.yaml, but it cannot write rules_engine.py or mandates.",
        "Every fix is simulated against the rules engine before a human approves it.",
        "Shadow state (effective portfolios) records approved changes; the seed stays immutable.",
    ])

    # 3. Architecture
    add_heading(doc, "3. Architecture", level=1)
    add_para(doc, "Next.js 14 App Router frontend + FastAPI backend, monorepo.")
    add_code(doc, [
        "[Next.js: /, /portfolio/{id}, /portfolio/{id}/workbench, /hermes]  ←HTTP→  [FastAPI routers]",
        "                                        │                              ├→ agents/: explain, remediate, summarize",
        "                                        │                              ├→ agents/hermes/: proposer, loop, score, reflect",
        "                                        │                              ├→ core/: rules_engine (LAW), effective (shadow state)",
        "                                        │                              └→ core/hermes_store.py (runtime state)",
        "                                        ↓                                    ↑ grounded against",
        "                          [rules_engine.py]  ← deterministic source of truth",
        "                                        ↓",
        "                [34,000 synthetic portfolios + mandates]  ← SQLite WAL + seeded GBM market model",
        "                                        ↓",
        "                          [audit.jsonl — append-only]",
    ])
    add_heading(doc, "Key backend modules", level=2)
    add_bullets(doc, [
        "backend/core/rules_engine.py — pure-function mandate checks; returns {status, breaches, watches, per_rule}.",
        "backend/core/effective.py — applies approved trades to seed portfolios (shadow state).",
        "backend/core/data_loader.py — SQLite data access with thread-local connections for Uvicorn worker safety.",
        "backend/core/hermes_store.py — pluggable store (SQLite local / Postgres production) for scan jobs, heartbeat and queue.",
        "backend/agents/hermes/loop.py — book-wide scan, proposer gating, queue ranking, heartbeat writes.",
        "backend/core/market.py — seeded GBM price evolution and virtual clock endpoints.",
    ])

    # 4. Pages & click-path
    add_heading(doc, "4. Page-by-Page Demo Script", level=1)

    add_heading(doc, "4.1 Command Centre (/)", level=2)
    add_para(doc,
        "The landing page shows a FUM-sized treemap. Tiles are sorted by FUM (largest top-left), so the first "
        "screen always shows a mix of red, orange and green. The summary bar narrates the whole book in plain English.")
    add_bullets(doc, [
        "Red = breached portfolios. Orange = drifting past target + tolerance. Green = aligned.",
        "The status reflects the effective (post-approval) shadow state, not the stale seed.",
        "Market Panel is visible at the top so the presenter can tick time forward and show live drift.",
    ])

    add_heading(doc, "4.2 Market Panel", level=2)
    add_para(doc,
        "A seeded GBM model drives all 34,000 portfolios. Click Tick or Advance N Days to move the virtual clock. "
        "Prices, market values and status are recalculated. API errors are surfaced in the panel instead of failing silently.")
    add_bullets(doc, [
        "POST /market/tick — advance one day.",
        "POST /market/advance?days=N — batch advance.",
        "POST /market/auto-run — toggle background ticking.",
    ])

    add_heading(doc, "4.3 Diagnosis (/portfolio/{id})", level=2)
    add_para(doc,
        "The narrative panel translates the deterministic rules result into plain English. Below it, breach chips "
        "list every rule the engine flagged. Click a chip to highlight the offending holdings in the table. "
        "Click the ? button on a row to explain the exact rule that applies to that holding or asset class.")
    add_bullets(doc, [
        "Confidence line separates deterministic rule-maths from AI advisory.",
        "HoldingsTable uses the per_rule row that mentions the ticker, asset class, sector or region.",
        "AllocationBarChart targets the top asset-class bar with its matching rule.",
    ])

    add_heading(doc, "4.4 Workbench (/portfolio/{id}/workbench)", level=2)
    add_para(doc,
        "The human-in-the-loop remediation loop. AI proposes minimal compliant trades, the verification panel "
        "simulates them against the rules engine, and Approve writes the shadow state + audit trail.")
    add_bullets(doc, [
        "Propose a fix → POST /portfolio/{id}/remediate (AI-driven, rules-gated retry).",
        "Verify → POST /portfolio/{id}/verify runs rules_engine on the simulated portfolio.",
        "Approve → POST /portfolio/{id}/approve applies the trade via effective.py and re-checks.",
        "Export CSV → RFC-4180 quoted fields for Excel safety.",
    ])

    add_heading(doc, "4.5 Hermes Mission Control (/hermes)", level=2)
    add_para(doc,
        "Hermes scans the whole book, proposes compliant remediations, ranks them, and reflects on misses to "
        "suggest ONE strategy change at a time. Humans adopt or reject each suggestion.")
    add_bullets(doc, [
        "Scan Book → POST /hermes/scan returns a job id; GET /hermes/scan/{job_id} polls progress.",
        "Queue → GET /hermes/queue lists ranked, rules-gated proposals (local run: 938 verified, 762 missed, 5,100 drift-only skips out of 34,000).",
        "Approve Batch → POST /hermes/approve-batch applies selected rows through the same human gate.",
        "Reflect → POST /hermes/reflect proposes a strategy tweak.",
        "Adopt → POST /hermes/adopt writes the new strategy version and archives the old one.",
        "History → GET /hermes/history shows every versioned change.",
    ])
    add_para(doc,
        "A recent hardening pass fixed the Scan Book engine so it no longer returns zero verified rows. "
        "The proposer now handles all ten breach types, chooses mandate-safe replacements, respects top-N concentration, "
        "and reserves a trade slot for the redeploy buy. Misses that remain are genuinely infeasible under the current mandate.",
        italic=True)

    # 5. API reference
    add_heading(doc, "5. API Reference (FastAPI)", level=1)
    add_para(doc, "All routes are relative to the backend root and proxied from /api/* in production.")
    api_groups = [
        ("Portfolios", [
            "GET /portfolios",
            "GET /portfolio/{client_id}",
            "GET /portfolio/{client_id}/check",
            "GET /portfolios/summary",
            "GET /portfolios/summary_ai",
            "GET /portfolios/top",
        ]),
        ("Actions", [
            "POST /portfolio/{client_id}/explain",
            "POST /portfolio/{client_id}/verify",
            "POST /portfolio/{client_id}/remediate",
            "POST /portfolio/{client_id}/approve",
            "POST /portfolio/{client_id}/reflect",
            "POST /preferences/adopt",
        ]),
        ("Market", [
            "GET /market/status",
            "GET /market/clock",
            "POST /market/tick",
            "POST /market/advance",
            "POST /market/auto-run",
            "POST /market/auto-fix",
            "GET /market/prices",
            "GET /market/history",
        ]),
        ("Hermes", [
            "POST /hermes/scan",
            "GET /hermes/scan/{job_id}",
            "GET /hermes/queue",
            "POST /hermes/approve-batch",
            "GET /hermes/strategy",
            "POST /hermes/reflect",
            "POST /hermes/adopt",
            "GET /hermes/heartbeat",
            "GET /hermes/history",
            "POST /hermes/rollback",
        ]),
        ("Admin / Audit", [
            "POST /admin/reset",
            "GET /audit",
            "GET /health",
        ]),
    ]
    for title, routes in api_groups:
        add_heading(doc, title, level=2)
        add_bullets(doc, routes)

    # 6. Quality assurance
    add_heading(doc, "6. Quality Assurance & Tests", level=1)
    add_para(doc, "The project ships with unit, component, integration and UI validation tests.")
    add_code(doc, [
        "cd backend && .venv/Scripts/python.exe -m pytest           # backend unit + Hermes store + scaled tests",
        "cd frontend && npx tsc --noEmit                            # TypeScript type check",
        "cd frontend && npx next build                              # static production build",
        "cd frontend && npx vitest run                              # React component tests",
        "python scripts/e2e_screenshots.py                          # Playwright UI smoke test + screenshots",
    ])
    add_bullets(doc, [
        "Backend pytest: 143 passed, 1 skipped.",
        "Frontend: tsc --noEmit clean, next build succeeds, vitest 6 passed.",
        "Recent hardening + UI work: thread-local SQLite connections (concurrent Uvicorn workers), isolated Hermes store "
        "connection, pluggable HermesStore with Postgres backend for serverless deploys, RFC-4180 CSV export, "
        "targeted explain metrics, MarketPanel error surfacing, Playwright screenshot validation, and a full "
        "Muted Light Matrix UI makeover across Command Centre, Diagnosis, Workbench and /hermes.",
    ])

    # 7. Deployment
    add_heading(doc, "7. Deployment", level=1)
    add_para(doc,
        "Backend is configured for Render via backend/render.yaml. Frontend is configured for Vercel via vercel.json. "
        "Set ANTHROPIC_API_KEY as a server-side environment variable on Render. The frontend uses a server-only "
        "API_URL env and next.config rewrites so the backend URL is never baked into the client bundle.")
    add_bullets(doc, [
        "Render: builds the backend, installs requirements, starts uvicorn.",
        "Vercel: builds the Next.js app; /api/* rewrites to API_URL.",
        "Serverless SQLite caveat solved by hermes_store.py Postgres backend for production.",
    ])

    # 8. Runbook
    add_heading(doc, "8. Demo Runbook for Kevin", level=1)
    add_para(doc, "A 3-minute click-path that hits every headline feature.")
    steps = [
        "Open http://localhost:3000. Read the summary bar: 34,000 portfolios, counts of green/orange/red.",
        "Tick the Market Panel once or twice; watch the summary counts shift as prices drift.",
        "Click a large red tile → Diagnosis. Read the AI narrative, then click the top breach chip to highlight offenders.",
        "Click ? on an offending holding. The popover names the exact rule and its current/limit values.",
        "Click Workbench. Press Propose a fix, inspect the trades, press Verify, then Approve. Watch status flip green.",
        "Click Export CSV and open the file in Excel to show quoted fields survive commas/whitespace.",
        "Navigate to /hermes. Press Scan Book, wait for the queue, inspect the Book Score and Strategy panels.",
        "Press Reflect, read the proposed strategy change, press Adopt. Check History for the archived version.",
        "Select a few queue rows and press Approve Batch; observe the queue marks them processed.",
        "Reset the book from /admin/reset if you want to re-run from a clean state.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f"{i}. {step}", style="List Number")
        p.paragraph_format.left_indent = Inches(0.25)

    # 9. Notes / support
    add_heading(doc, "9. Notes & Support", level=1)
    add_para(doc,
        "If something looks wrong, check backend/backend.log and frontend/frontend.log. "
        "Common gotchas: the 34,000-portfolio SQLite file is created on first startup by main.py's startup handler "
        "or by running generators/generate_data.py manually. Without ANTHROPIC_API_KEY all LLM calls return deterministic "
        "MockLLM text so the demo works offline.")

    # Save
    doc.save(DOCX_PATH)
    print(f"Wrote {DOCX_PATH}")


if __name__ == "__main__":
    main()
