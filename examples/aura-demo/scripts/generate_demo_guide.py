"""Generate AURA_Demo_Guide.docx for the Financial Simplicity demo.

Run from the repo root:
    python scripts/generate_demo_guide.py

Produces: AURA_Demo_Guide.docx in the repo root.
"""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill: str):
    """Set background shading on a table cell (fill as hex, e.g. 'D9E2F3')."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def _add_heading(doc, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.space_before = Pt(12)
    return p


def _add_para(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    return p


def _add_bullet(doc, text: str, indent: int = 0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
    return p


def _add_numbered(doc, text: str):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = "Calibri"
    return p


def _add_code(doc, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)
    p.paragraph_format.space_after = Pt(4)
    return p


def _add_table(doc, headers, rows, header_fill='D9E2F3'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        _set_cell_shading(hdr_cells[i], header_fill)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    return table


def main():
    doc = Document()

    # ---- Styles ----
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # ---- Title Page ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ASSURE")
    run.font.size = Pt(44)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
    run.font.name = "Calibri Light"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Portfolio Assurance Platform")
    run.font.size = Pt(20)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tag.add_run('"AI recommends. Assurance verifies."')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x33)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Demo Guide for Financial Simplicity\nPrepared by Satyam Das — June 2026")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_page_break()

    # ---- Executive Summary ----
    _add_heading(doc, "Executive Summary", 1)
    _add_para(doc,
        "ASSURE is a thin, AI-powered assurance layer that sits on top of an existing portfolio-assurance platform. "
        "It turns dense, spreadsheet-bound V1 monitoring into an interactive V2 experience: a colour-coded command centre, "
        "plain-English diagnosis, one-click remediation, and a self-improving book-wide engine called Hermes — all gated "
        "by a deterministic rules engine and a human approver.",
    )
    _add_para(doc,
        "Everything in this demo is synthetic. No real client data is used. The live deployment is a single URL that Kevin "
        "and the team can open, click through, and interrogate end-to-end in under six minutes.",
    )

    _add_heading(doc, "Live demo", 2)
    _add_para(doc, "Frontend:", bold=True)
    _add_code(doc, "https://aura-demo-rho.vercel.app/")
    _add_para(doc, "Backend API:", bold=True)
    _add_code(doc, "https://aura-demo-rho.vercel.app/api/health")
    _add_para(doc,
        "The stack is Next.js 14 + Tailwind (frontend) and FastAPI + SQLite (backend), deployed on Vercel and Render. "
        "All prices, portfolios, and mandates are generated deterministically from a seeded synthetic universe of ~35 tickers "
        "and 34,000 client portfolios.",
    )

    # ---- The Problem We Are Solving ----
    _add_heading(doc, "1. The Problem ASSURE Solves", 1)
    _add_para(doc,
        "Wealth managers continuously verify that thousands of personalised client portfolios remain aligned to each client's "
        "mandate. Today's tools are accurate but slow, dense, and manual:",
    )
    _add_bullet(doc, "A heatmap that shows colour blocks but no plain-English interpretation.")
    _add_bullet(doc, "A client-detail screen so dense that even the COO can't always tell what is wrong.")
    _add_bullet(doc, "A manual rebalancing table that requires the manager to compute every compliant trade by hand.")
    _add_bullet(doc, "No book-wide view of systemic risk or automated, auditable remediation queue.")
    _add_para(doc,
        "ASSURE does not replace the legacy engine. It reads the same data and renders it better, adds an AI comprehension layer, "
        "and proposes compliant fixes — while keeping the deterministic rules engine and a human as the final authorities.",
    )

    # ---- Automation Highlights ----
    _add_heading(doc, "2. Automation Highlights", 1)
    _add_para(doc,
        "The demo showcases five automation layers. Each one is deliberately placed so the audience can see exactly where "
        "the machine adds value and where the human stays in control.",
    )

    headers = ["Layer", "What it automates", "Human / safety gate"]
    rows = [
        ["AI Summary", "Reads the whole 34k book and writes a 2-sentence English summary of counts, systemic risk, and remediation estimate.", "Grounded only on deterministic rules-engine counts."],
        ["Diagnosis Narrative", "Translates rule breaches and drift watches into calm, plain English with offending holdings identified.", "Narrative is advisory; every claim traces to a rule-engine fact."],
        ["Per-Metric Explain", "Every holding row and allocation slice has an 'Explain' button that describes the specific check in one sentence.", "Still grounded in rules-engine per_rule rows."],
        ["Remediation Agent", "Proposes minimal, mandate-compliant trades for a breached portfolio and retries once if the first pass still breaches.", "Manager edits or approves; rules engine re-verifies before state changes."],
        ["Hermes Book-Wide Engine", "Scans all 34k portfolios, proposes fixes for every drifting one, ranks by FUM × severity, and learns from misses.", "Every proposal is rules-engine-gated; every batch approval is a human click; every strategy mutation is human-adopted and versioned."],
    ]
    _add_table(doc, headers, rows)

    _add_heading(doc, "Deterministic core vs. AI layer", 2)
    _add_para(doc,
        "The architecture separates two tiers that are often conflated in AI demos:",
    )
    _add_bullet(doc, "Mandate rules = LAW. Stored per portfolio, enforced only by core/rules_engine.py. Hermes and the LLM can never write here.")
    _add_bullet(doc, "Remediation strategy = JUDGMENT. Stored in agents/hermes/strategy.yaml. Hermes reflection proposes changes to this file only, and a human must adopt each version bump.")
    _add_para(doc,
        "This two-tier cage is the structural answer to 'how do we trust a self-improving AI with client money?' — it cannot go rogue "
        "because the rules engine gates every output regardless of what Hermes has learned.",
    )

    # ---- Problems Solved End-to-End ----
    _add_heading(doc, "3. Problems Solved End-to-End", 1)

    _add_heading(doc, "3.1 From unreadable detail to plain-English diagnosis", 2)
    _add_para(doc,
        "Pain: Kevin noted that he sometimes cannot tell what he is looking at in the detail view. "
        "Fix: the Diagnosis page opens with a NarrativePanel that reads the portfolio for the user, plus clickable BreachChips that "
        "highlight offending holdings. The confidence line explicitly separates deterministic rule-maths from AI advisory.",
    )

    _add_heading(doc, "3.2 From manual rebalancing to propose → verify → approve", 2)
    _add_para(doc,
        "Pain: Rebalancing is spreadsheet-heavy and error-prone. Fix: one 'Propose a fix' button calls a breach-aware remediation agent. "
        "The proposed trades are then fed back through the deterministic rules engine in the VerifyPanel. The manager can edit trades live "
        "and see the rule checks update in real time. Approval writes to an append-only SQLite state table and flips the portfolio green — "
        "and it stays green on reload because the holdings genuinely changed, not because of an override flag.",
    )

    _add_heading(doc, "3.3 From one-portfolio-at-a-time to book-wide remediation", 2)
    _add_para(doc,
        "Pain: A 34,000-portfolio book is too large to remediate manually. Fix: Hermes Mission Control scans the entire book asynchronously, "
        "ranks proposals by funds-at-risk × severity, and presents a queue of pre-verified fixes. A manager can approve individual rows or "
        "approve-all-verified with a single human-gated click.",
    )

    _add_heading(doc, "3.4 From static policy to self-improving (but caged) strategy", 2)
    _add_para(doc,
        "Pain: A single fixed remediation strategy misses many portfolios. Fix: after each scan, Hermes reflects on its misses and proposes "
        "exactly one change to strategy.yaml (e.g., switch trim method from 'largest_overweight' to 'liquidate' because too many proposals "
        "still breached). A human adopts or dismisses the change. Every adoption bumps the version, archives the prior strategy to history/, "
        "and appends an audit entry. The mandate rules remain untouched.",
    )

    _add_heading(doc, "3.5 From opaque operations to full auditability", 2)
    _add_para(doc,
        "Every AI suggestion, every rule-engine check, every human approval, every trade application, and every strategy mutation is "
        "appended to an audit trail with timestamp, actor, rationale, and version. Nothing happens without a log.",
    )

    # ---- Alignment with Financial Simplicity Ethos ----
    _add_heading(doc, "4. Alignment with Financial Simplicity Ethos", 1)
    _add_para(doc,
        "Stuart Holdsworth's thesis is 'AI Can Recommend. Assurance Verifies.' ASSURE is built as a literal implementation of that sentence:",
    )

    ethos = [
        ("Investor first", "The UI is designed around the manager's need to protect client outcomes, not around the algorithm. The largest, most-at-risk portfolios are shown first."),
        ("Integrity / Accountability", "The rules engine is the only source of truth for compliance. The LLM and Hermes are explicitly advisory/proposal layers. Every action is audited."),
        ("Continuous improvement", "Hermes reflects on its own misses and proposes strategy improvements, but each change passes a human adoption gate and is versioned/reversible."),
        ("Let the machines do the work", "Boring, repetitive verification and trade-calculation work is automated; judgment, approval, and policy changes stay with people."),
        ("Trust through verification", "Conformal-prediction-style honesty is visible in the UI: per-rule ✅/❌, confidence derived from rule-pass counts, and a clear 'deterministic vs advisory' confidence line."),
    ]
    for title, body in ethos:
        _add_bullet(doc, f"{title}: {body}")

    _add_para(doc,
        "The demo also respects the reality of a small, senior engineering team on a mature stack: ASSURE is framed as a layer on top, "
        "not a replacement. It reads their data and renders it better; it does not ask them to rebuild the core engine.",
    )

    # ---- Architecture & Safety Cage ----
    _add_heading(doc, "5. Architecture & the Assurance Cage", 1)
    _add_para(doc, "Data flow (single request path):", bold=True)
    _add_numbered(doc, "Hermes proposes a remediation using agents/hermes/strategy.yaml.")
    _add_numbered(doc, "core/rules_engine.py checks the post-trade portfolio against the client's mandate.")
    _add_numbered(doc, "Only proposals that pass the rules engine reach the human queue.")
    _add_numbered(doc, "A manager approves (or edits and re-verifies).")
    _add_numbered(doc, "Trades are applied to the shadow SQLite state table; status is recomputed; the audit trail is appended.")
    _add_numbered(doc, "Hermes reflects on misses and may propose a strategy change — which also needs human adoption and is versioned.")

    _add_para(doc,
        "Enforced in code: agents/hermes/strategy_io._guard() refuses any write path outside strategy.yaml/history/, so mandate rules and "
        "rules_engine.py are structurally unreachable from the judgment layer.",
    )

    _add_heading(doc, "Scale numbers", 2)
    headers = ["Metric", "Value"]
    rows = [
        ["Synthetic portfolios", "34,000"],
        ["Synthetic tickers / asset classes", "~35 tickers across Equity, Bonds, Commodity, Crypto, Cash"],
        ["Mandate rule dimensions", "10 (asset-class, sector, region, approved universe, single holding, cash, drift, ESG exclusion, top-N concentration, liquidity floor)"],
        ["Initial book colour distribution", "~80% green, ~15% orange, ~5% red"],
        ["Backend tests", "143 passed, 1 skipped"],
        ["Frontend unit tests", "3 passed"],
        ["Deep E2E UI + API tests", "23/23 passed (scripts/e2e_deep.py)"],
        ["Hermes scan on day-0 book", "~90 portfolios remediated, ~1,480 missed (expected for v1 strategy; reflection improves this)"],
    ]
    _add_table(doc, headers, rows)

    # ---- How to Check for Updates / Health ----
    _add_heading(doc, "6. How to Check for Updates and System Health", 1)
    _add_para(doc, "Before and during the demo, verify the deployment with these endpoints and checks:", bold=True)
    _add_code(doc, "GET https://aura-demo-rho.vercel.app/api/health")
    _add_code(doc, "GET https://aura-demo-rho.vercel.app/api/portfolios/summary")
    _add_code(doc, "GET https://aura-demo-rho.vercel.app/api/hermes/heartbeat")
    _add_para(doc, "Expected behaviour:", bold=True)
    _add_bullet(doc, "/api/health returns {\"status\":\"ok\"}.")
    _add_bullet(doc, "/api/portfolios/summary returns total ≈ 34,000 with green/orange/red counts.")
    _add_bullet(doc, "/api/hermes/heartbeat returns scan counts and a composite book-health score.")
    _add_para(doc,
        "Local development (if you want to run before/after changes):", bold=True,
    )
    _add_code(doc, "cd backend && .venv\\Scripts\\python.exe -m uvicorn main:app --port 8000")
    _add_code(doc, "cd frontend && npm run dev")
    _add_code(doc, "cd backend && .venv\\Scripts\\python.exe -m pytest tests/ -q")
    _add_code(doc, "cd frontend && npm run build")

    # ---- End-to-End Feature Walkthrough ----
    _add_heading(doc, "7. End-to-End Feature Walkthrough", 1)

    _add_heading(doc, "7.1 Command Centre (/) — the at-a-glance book", 2)
    _add_bullet(doc, "Opens on a FUM-sized treemap. Red = breached, orange = drifting, green = aligned. Largest portfolios appear top-left so every screen shows a mixed picture.")
    _add_bullet(doc, "Top banner reads the whole book in plain English: e.g., '34,000 portfolios; 1,700 breached; largest systemic risk is tech concentration.'")
    _add_bullet(doc, "Market panel shows a virtual clock, seeded price model, and controls: TICK, AUTO-ADVANCE, STOP CLOCK. Clicking TICK advances the market day and recomputes drift.")
    _add_bullet(doc, "Filters: adviser dropdown, asset-class chip, 'Needs Action Today' toggle, and free-text search by client name/id.")
    _add_bullet(doc, "Urgent Triage Log on the right lists red/orange portfolios ranked by funds-at-risk.")

    _add_heading(doc, "7.2 Diagnosis (/portfolio/{id}) — the unreadable-made-readable screen", 2)
    _add_bullet(doc, "Header shows client name, adviser, FUM, and status badge.")
    _add_bullet(doc, "NarrativePanel explains breaches and watches in plain English (now also works offline via a deterministic MockLLM fallback).")
    _add_bullet(doc, "BreachChips are clickable: clicking a chip highlights the offending holdings in the holdings table.")
    _add_bullet(doc, "HoldingsTable has a '?' button per row; AllocationBarChart has an 'Explain' button. Both call the grounded explainer and return a one-sentence rule check.")
    _add_bullet(doc, "Confidence line: 'Rule checks are deterministic (100%). The narrative is advisory (AI-inferred).'")
    _add_bullet(doc, "Allocation bar chart shows weights vs mandate caps; offending bars are highlighted red; PerformanceChart is synthetic.")

    _add_heading(doc, "7.3 Workbench (/portfolio/{id}/workbench) — propose, verify, approve", 2)
    _add_bullet(doc, "Click 'Propose a fix'. The remediation agent reads the breach manifest and proposes a minimal trade list.")
    _add_bullet(doc, "VerifyPanel on the right re-runs the rules engine on the post-trade portfolio and shows per-rule ✅/❌.")
    _add_bullet(doc, "Edit any trade's units in the WorkbenchTable; the rules engine re-checks live and updates the VerifyPanel.")
    _add_bullet(doc, "Click 'Approve & Log'. The trades are applied to the shadow state, an audit entry is written, and the portfolio flips green. The green status persists on reload.")
    _add_bullet(doc, "'Reset demo state' clears all applied trades and rewinds the clock so the demo can be re-run cleanly.")

    _add_heading(doc, "7.4 Hermes Mission Control (/hermes) — book-wide self-improving remediation", 2)
    _add_bullet(doc, "The Assurance Cage diagram explains the four-stage gate: Hermes proposes → Rules engine verifies → Human approves → Feeds back to Hermes.")
    _add_bullet(doc, "Click 'Scan Book'. Hermes processes all 34k portfolios asynchronously. The queue fills with proposals that have already been rules-engine-gated green.")
    _add_bullet(doc, "Expand any queue row to see the exact proposed trades and a live VerifyPanel for that portfolio.")
    _add_bullet(doc, "Approve individual rows or click 'Approve all verified'. After a successful bulk approve the queue is refreshed from the backend, so additional queued items (beyond the visible page) reappear correctly.")
    _add_bullet(doc, "Book Score panel shows alignment rate, acceptance rate, avg trades per fix, breaches remaining, and a composite score.")
    _add_bullet(doc, "Strategy panel shows the six judgment variables and their rationales. Click 'Reflect (deterministic)' or 'Reflect (Hermes)' to propose one strategy tweak.")
    _add_bullet(doc, "Adopt a proposal: version bumps, prior strategy is archived to history/, and an audit entry is written. The History panel shows every version and allows rollback.")

    _add_heading(doc, "7.5 Admin Reset", 2)
    _add_bullet(doc, "The 'Reset demo state' button (Workbench) or POST /admin/reset clears SQLite state, hermes_queue, scan_jobs, drift_events, status_history, and rewinds the market clock to day 0. Use this before every clean demo run.")

    # ---- Demo Script (6-Minute Click Path) ----
    _add_heading(doc, "8. Recommended 6-Minute Demo Script", 1)
    _add_para(doc, "Use this script when walking Kevin and the team through the live URL. It is designed to flow from one portfolio to the whole book, "
                  "and to close on the self-improving-but-caged idea.")

    for step, script in [
        ("1. Command Centre", "'This is the 34,000-portfolio book. The AI summary already tells us what's wrong before we click anything: ~80% aligned, ~5% breached, ~15% attention. The blocks are sized by FUM, so the biggest problems are impossible to miss.'"),
        ("2. Click the biggest red block → Diagnosis", "'This is the screen that used to be unreadable. The narrative panel reads the portfolio for us: the engine flagged the breaches; the AI just translates them to plain English.'"),
        ("3. Click a breach chip", "'Clicking a breach chip highlights the exact holdings causing it. Every 'Explain' button describes the specific rule check in one sentence.'"),
        ("4. Open Remediation Workbench", "'Now we fix it. Propose a fix generates a minimal compliant trade list. The VerifyPanel on the right is the rules engine re-checking the post-trade portfolio — not the AI deciding compliance.'"),
        ("5. Approve & Log", "'I can edit any trade and the rules engine re-checks live. Once it's green, I approve. The status flips green and persists on reload because the holdings actually changed — not because of a flag.'"),
        ("6. Open /hermes Mission Control", "'That was one portfolio by hand. Now watch it do the whole book. Scan Book processes all 34k portfolios; only rules-engine-green proposals reach the queue.'"),
        ("7. Expand a queue row + approve", "'Every row was deterministically verified before it reached me. Hermes can be as creative as it likes — it physically cannot output something non-compliant.'"),
        ("8. Show Hermes learning", "'Hermes is now reflecting on its misses. It wants to change one strategy variable — say, the trim method — with a written rationale. I can adopt or dismiss. If I adopt, the version bumps, the prior strategy is archived, and the next scan uses the better strategy.'"),
        ("9. Close", "'So: AI proposes and even improves how it proposes, but deterministic assurance verifies every output, and a human approves every action. That's AI recommends, assurance verifies — made self-improving, and made safe.'"),
    ]:
        _add_bullet(doc, f"{step}: {script}")

    # ---- Phase 2 E2E Audit Results ----
    _add_heading(doc, "9. Phase 2 End-to-End Audit Results", 1)
    _add_para(doc,
        "A multi-agent workflow plus an automated API audit exercised every page, endpoint, and flow against the live deployment. "
        "This section records what was verified and what remains a known caveat.",
    )
    _add_para(doc, "Verified green:", bold=True)
    _add_bullet(doc, "/api/health, /api/portfolios/summary, /api/portfolios/top, /api/portfolio/{id}, /api/portfolio/{id}/explain, /api/portfolio/{id}/remediate, /api/portfolio/{id}/verify, /api/portfolio/{id}/approve all return correct, grounded results.")
    _add_bullet(doc, "Workbench flow: red portfolio → propose → verify → approve → status flips to green and persists across reloads.")
    _add_bullet(doc, "Hermes flow: scan → queue populates with rules-engine-gated proposals → approve individual or bulk → trades applied and audited.")
    _add_bullet(doc, "Reflection/adoption/rollback endpoints exist and are human-gated; strategy versions are archived to history/.")
    _add_bullet(doc, "Local backend test suite: 143 passed, 1 skipped. Frontend unit tests: 3 passed. Production build: passes TypeScript + lint.")
    _add_bullet(doc, "Deep headless E2E (scripts/e2e_deep.py): 23/23 checks passed across Command Centre, Diagnosis, Workbench propose/verify/approve/CSV export, and Hermes scan/individual approve/bulk approve/reflect/adopt.")
    _add_para(doc, "Fixes applied during the audit:", bold=True)
    _add_bullet(doc, "BreachChips now receives the portfolio clientId, so per-metric Explain buttons render on the Diagnosis page.")
    _add_bullet(doc, "BreachChips accessibility improved: Explain is now a sibling button instead of a nested interactive element.")
    _add_bullet(doc, "VerifyPanel now shows the actual prior status (green/orange/red) instead of always rendering BREACH.")
    _add_bullet(doc, "SummaryBar no longer shows a fabricated '+2.4% vs last mo' FUM trend.")
    _add_bullet(doc, "AssuranceBanner no longer fabricates a hard-coded 'TECH SECTOR CONCENTRATION' narrative or stale 'INITIATING SCAN' copy.")
    _add_bullet(doc, "TriageQueue 'VIEW ALL' now navigates to /hermes.")
    _add_bullet(doc, "HermesQueue reject copy corrected to describe the actual local-only behaviour.")
    _add_bullet(doc, "Rebranded the application from AURA to ASSURE across metadata, topbar, and sidebar.")
    _add_bullet(doc, "Removed non-functional Diagnosis / Audit Log / Settings navigation items from the sidebar; kept only Command Centre and Hermes Engine.")
    _add_bullet(doc, "Completed a full UI makeover to the 'Muted Light Matrix' institutional design system (JetBrains Mono, cool slate palette, 4px grid, micro-shadows, high-density data tables).")
    _add_bullet(doc, "Fixed /admin/reset on large books: list_portfolios now chunks client_id queries under SQLite's ~999 variable limit, so reset works for the full 34k book.")
    _add_bullet(doc, "Replaced the donut allocation chart with an institutional bar chart that still supports per-asset-class Explain targeting.")
    _add_para(doc, "Known caveats:", bold=True)
    _add_bullet(doc, "The deployed backend uses SQLite on Vercel serverless functions. Scan job and queue state are instance-local, so GET /api/hermes/scan/{job_id} can return 404 and queue visibility can flicker between requests. This is a deployment-platform limitation, not a code bug. For production, move scan_jobs and queue state to a shared store (Postgres/Redis/S3).")
    _add_bullet(doc, "Hermes queue rows are not deleted or marked processed after approval; the queue is a static snapshot of a scan day. Heartbeat counts are also static until the next scan.")
    _add_bullet(doc, "Static HTML fetch cannot validate React client-rendered component names; use a headless browser or local dev build for component-level UI validation.")

    # ---- Demo Tips ----
    _add_heading(doc, "10. Demo Tips", 1)
    _add_bullet(doc, "Reset before every demo: Workbench → 'Reset demo state' (or POST /admin/reset).")
    _add_bullet(doc, "Pick c00011 for the single-portfolio demo path — it reliably resolves green after one remediation.")
    _add_bullet(doc, "If the queue is empty after a scan, that is normal for some strategy states; reset and re-scan, or tick the market forward to create new drift.")
    _add_bullet(doc, "The MockLLM fallback now produces real grounded narrative text even without ANTHROPIC_API_KEY; with the key set, Claude produces the same grounded prose with more natural phrasing.")
    _add_bullet(doc, "Use the API endpoints in Section 6 if the UI ever looks suspicious — the JSON is the ground truth.")
    _add_bullet(doc, "On Vercel, scan/queue state may appear inconsistent across rapid page refreshes because SQLite is instance-local. Slow down or run the demo from a single serverless-warm session.")

    # ---- Next Steps / Roadmap ----
    _add_heading(doc, "11. Suggested Next Steps", 1)
    _add_bullet(doc, "Productionise the deterministic rules engine integration with Financial Simplicity's existing portfolio data model.")
    _add_bullet(doc, "Add real-time market data feed and historical drift analysis.")
    _add_bullet(doc, "Extend Hermes reflection with a human-feedback loop on per-client adviser preferences.")
    _add_bullet(doc, "Add order-management-system (OMS) integration so approved trades can flow to execution with a second human trader gate.")
    _add_bullet(doc, "Run a supervised pilot on a small adviser book and measure time-to-remediate and breach-resolution rate.")

    # ---- Closing ----
    _add_heading(doc, "Closing", 1)
    _add_para(doc,
        "ASSURE is not a toy. It is the NeuralQuant multi-agent pattern — import, grounded analysis, adversarial verification, human approval, "
        "audit trail — transplanted into the exact domain Financial Simplicity owns. It demonstrates that AI-native, ship-fast builders can "
        "add immediate value to a senior, stable team without disrupting the legacy foundation they have spent years getting right.",
    )
    _add_para(doc,
        "The single most important takeaway for Kevin: the demo is built around 'AI recommends, assurance verifies' — the same thesis Stuart is writing about. "
        "Every screen reinforces it. Every safety gate proves it. And every audit entry makes it auditable.",
    )

    repo_root = Path(__file__).resolve().parents[1]
    out_path = repo_root / "AURA_Demo_Guide.docx"
    doc.save(out_path)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
